#!/usr/bin/env python3
"""
Format a Word (.docx) report to a professional style.

Applies:
- 1" margins
- Normal body: Times New Roman 12pt, justified, 1.15 line spacing, 6pt after
- Headings (1-3): bold, consistent sizes, spacing, alignment
- Footer with centered page number (Page X of Y)
- Optional header with document title
- Optional Table of Contents placeholder (Word will populate on Update Field)
- Bulleted/numbered list normalization

Usage:
  python scripts/format_report.py --input "/path/to/NYC_Taxi_Fare_Report.docx" \
      --output "/path/to/NYC_Taxi_Fare_Report_formatted.docx" \
      --title "NYC Taxi Fare Prediction Report" --toc

Note: Word may require you to open the output and press: Select All (Cmd+A) → F9 (Update Fields)
      to render TOC and page numbers correctly.
"""
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

HEADING_SPECS = {
    1: {"size": 16, "bold": True, "space_before": 12, "space_after": 6, "align": WD_ALIGN_PARAGRAPH.LEFT},
    2: {"size": 14, "bold": True, "space_before": 10, "space_after": 4, "align": WD_ALIGN_PARAGRAPH.LEFT},
    3: {"size": 12, "bold": True, "space_before": 8, "space_after": 2, "align": WD_ALIGN_PARAGRAPH.LEFT},
}

COMMON_HEADING_KEYWORDS = {
    1: ["abstract", "introduction", "overview", "background", "conclusion", "appendix", "references"],
    2: ["dataset", "data", "method", "methodology", "experiments", "results", "discussion", "limitations", "future work"],
}


def set_margins(section, inches=1.0):
    section.top_margin = Inches(inches)
    section.bottom_margin = Inches(inches)
    section.left_margin = Inches(inches)
    section.right_margin = Inches(inches)


def ensure_styles(doc: Document):
    styles = doc.styles
    # Normal
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    # Headings 1-3
    for level in (1, 2, 3):
        style_name = f'Heading {level}'
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        st = styles[style_name]
        st.font.name = 'Times New Roman'
        st.font.size = Pt(HEADING_SPECS[level]["size"])
        st.font.bold = HEADING_SPECS[level]["bold"]


def format_paragraph(paragraph, is_heading_level=None):
    pfmt = paragraph.paragraph_format
    if is_heading_level:
        spec = HEADING_SPECS[is_heading_level]
        paragraph.alignment = spec["align"]
        pfmt.space_before = Pt(spec["space_before"])
        pfmt.space_after = Pt(spec["space_after"])
    else:
        # Body text defaults
        paragraph.style = 'Normal'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pfmt.line_spacing = 1.15
        pfmt.space_after = Pt(6)


def detect_heading_level(text: str):
    t = text.strip()
    if not t:
        return None
    tl = t.lower()

    # Numbered heading patterns like "1.", "1.1", etc.
    if tl[:1].isdigit():
        return 1 if tl.count('.') <= 1 else 2

    # Keyword heuristics
    for level, words in COMMON_HEADING_KEYWORDS.items():
        if any(tl.startswith(w) for w in words):
            return level
    return None


def add_header_footer(doc: Document, title: str | None):
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # Header title (optional)
    if title:
        ph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        ph.text = title
        ph.style = doc.styles['Normal']
        ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ph.runs[0].font.bold = True

    # Footer: centered "Page X of Y"
    pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build: "Page " + PAGE + " of " + NUMPAGES
    pf.add_run("Page ")
    add_page_field(pf)
    pf.add_run(" of ")
    add_numpages_field(pf)


def add_field_run(paragraph, field_type: str):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), field_type)
    r = OxmlElement('w:r')
    fldSimple.append(r)
    t = OxmlElement('w:t')
    t.text = " "
    r.append(t)
    paragraph._p.append(fldSimple)


def add_page_field(paragraph):
    add_field_run(paragraph, 'PAGE')


def add_numpages_field(paragraph):
    add_field_run(paragraph, 'NUMPAGES')


def insert_toc(doc: Document):
    # Insert a Table of Contents field (Word must update fields to render)
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \o "1-3" \h \z \u')
    run._r.addnext(fld)


def normalize_lists(paragraph):
    # Simple heuristic: lines starting with '-', '•', or numbered '1)' -> convert to list appearance
    txt = paragraph.text.strip()
    if not txt:
        return
    if txt.startswith(('-', '•')):
        paragraph.style = 'List Bullet'
    elif (len(txt) > 2 and txt[0].isdigit() and txt[1] in '.)'):
        paragraph.style = 'List Number'


def process_document(input_path: str, output_path: str, title: str | None, add_toc: bool):
    doc = Document(input_path)

    # Margins
    for section in doc.sections:
        set_margins(section, 1.0)

    ensure_styles(doc)

    # Optional TOC at top (will require Word update fields)
    if add_toc:
        toc_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        toc_para.insert_paragraph_before("Table of Contents")
        toc_para.style = doc.styles['Heading 1']
        # real TOC field
        insert_toc(doc)
        doc.add_paragraph().add_run().add_break()

    # Header/Footer
    add_header_footer(doc, title)

    # Iterate paragraphs and apply formatting
    for p in doc.paragraphs:
        txt = (p.text or '').strip()
        level = detect_heading_level(txt)
        if level in (1, 2, 3):
            p.style = f'Heading {level}'
            format_paragraph(p, is_heading_level=level)
        else:
            format_paragraph(p)
            normalize_lists(p)

    doc.save(output_path)


def main():
    ap = argparse.ArgumentParser(description="Format a DOCX to professional style")
    ap.add_argument('--input', required=True, help='Path to input .docx')
    ap.add_argument('--output', required=True, help='Path to output .docx')
    ap.add_argument('--title', default=None, help='Document title to place in header')
    ap.add_argument('--toc', action='store_true', help='Insert a Table of Contents placeholder')
    args = ap.parse_args()

    process_document(args.input, args.output, args.title, args.toc)
    print(f"\n✓ Formatted document saved to: {args.output}")
    print("Open in Word and press Cmd+A then F9 to update fields (TOC/page numbers).")


if __name__ == '__main__':
    main()
